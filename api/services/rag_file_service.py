import csv
import re
import shutil
import subprocess
import tempfile
from collections.abc import Iterable, Mapping
from dataclasses import dataclass
from datetime import datetime, timezone
from functools import lru_cache
from pathlib import Path

import pandas as pd
import pdfplumber
from docx import Document as DocxDocument
from sqlalchemy import create_engine, text

from api.core.settings import PROJECT_ROOT, get_database_url
from utils.config_handler import chroma_conf
from utils.file_handler import get_file_sha256_hex


CREATE_RAG_FILES_SQL = """
CREATE TABLE IF NOT EXISTS rag_files (
    id BIGINT UNSIGNED NOT NULL AUTO_INCREMENT,
    original_name VARCHAR(255) NOT NULL,
    original_extension VARCHAR(16) NOT NULL,
    mime_type VARCHAR(255) NOT NULL DEFAULT '',
    size_bytes BIGINT UNSIGNED NOT NULL,
    sha256 CHAR(64) NOT NULL,
    source_path VARCHAR(512) NOT NULL,
    markdown_path VARCHAR(512) NOT NULL,
    markdown_content MEDIUMTEXT NOT NULL,
    status VARCHAR(16) NOT NULL DEFAULT 'pending',
    error_message TEXT NULL,
    chunk_count INT UNSIGNED NOT NULL DEFAULT 0,
    created_at DATETIME(6) NOT NULL DEFAULT CURRENT_TIMESTAMP(6),
    updated_at DATETIME(6) NOT NULL DEFAULT CURRENT_TIMESTAMP(6) ON UPDATE CURRENT_TIMESTAMP(6),
    PRIMARY KEY (id),
    UNIQUE KEY uq_rag_files_sha256 (sha256),
    KEY idx_rag_files_status_updated (status, updated_at),
    KEY idx_rag_files_name (original_name)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;
"""

ALLOWED_RAG_UPLOAD_EXTENSIONS = (
    ".doc",
    ".docx",
    ".txt",
    ".pdf",
    ".csv",
    ".md",
    ".xls",
    ".xlsx",
)

RAG_FILE_STATUS_PENDING = "pending"
RAG_FILE_STATUS_CONVERTING = "converting"
RAG_FILE_STATUS_INDEXING = "indexing"
RAG_FILE_STATUS_READY = "ready"
RAG_FILE_STATUS_FAILED = "failed"
RAG_FILE_INDEXABLE_STATUSES = (
    RAG_FILE_STATUS_READY,
    RAG_FILE_STATUS_INDEXING,
)


@dataclass(frozen=True)
class RagFileRecord:
    id: int
    original_name: str
    original_extension: str
    mime_type: str
    size_bytes: int
    sha256: str
    source_path: str
    markdown_path: str
    markdown_content: str
    status: str
    error_message: str | None
    chunk_count: int
    created_at: str
    updated_at: str


@dataclass(frozen=True)
class PreparedRagFile:
    record: RagFileRecord
    should_process: bool


def _now_utc() -> datetime:
    return datetime.now(timezone.utc).replace(tzinfo=None)


def _isoformat(value: object) -> str:
    if isinstance(value, datetime):
        normalized = value
    else:
        normalized = _now_utc()

    if normalized.tzinfo is None:
        normalized = normalized.replace(tzinfo=timezone.utc)

    return normalized.astimezone(timezone.utc).isoformat().replace("+00:00", "Z")


def _safe_original_name(filename: str) -> str:
    safe_name = Path(filename).name.strip()
    if not safe_name or safe_name in {".", ".."}:
        raise ValueError("文件名无效")

    if len(safe_name) > 255:
        raise ValueError("文件名不能超过 255 个字符")

    return safe_name


def _extension(filename: str) -> str:
    return Path(filename).suffix.lower()


def _relative_project_path(path: Path) -> str:
    return path.resolve().relative_to(PROJECT_ROOT.resolve()).as_posix()


def _resolve_project_path(path_value: str) -> Path:
    path = Path(path_value)
    if path.is_absolute():
        target = path.resolve()
    else:
        target = (PROJECT_ROOT / path).resolve()

    target.relative_to(PROJECT_ROOT.resolve())
    return target


def _clean_cell(value: object) -> str:
    text_value = str(value or "").replace("\r", " ").replace("\n", "<br>")
    return text_value.replace("|", "\\|").strip()


def _markdown_table(headers: Iterable[object], rows: Iterable[Iterable[object]]) -> str:
    header_values = [_clean_cell(header) or " " for header in headers]
    if not header_values:
        return ""

    lines = [
        "| " + " | ".join(header_values) + " |",
        "| " + " | ".join("---" for _ in header_values) + " |",
    ]

    for row in rows:
        cells = [_clean_cell(cell) for cell in row]
        if len(cells) < len(header_values):
            cells.extend("" for _ in range(len(header_values) - len(cells)))
        lines.append("| " + " | ".join(cells[: len(header_values)]) + " |")

    return "\n".join(lines)


def _dataframe_to_markdown(dataframe: pd.DataFrame) -> str:
    dataframe = dataframe.fillna("")
    return _markdown_table(dataframe.columns, dataframe.itertuples(index=False, name=None))


def _read_text(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue

    return path.read_text(encoding="utf-8", errors="ignore")


def _text_file_to_markdown(path: Path) -> str:
    return _read_text(path).strip()


def _csv_to_markdown(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            dataframe = pd.read_csv(path, dtype=str, keep_default_na=False, encoding=encoding)
            return _dataframe_to_markdown(dataframe)
        except UnicodeDecodeError:
            continue
        except Exception:
            break

    with path.open("r", encoding="utf-8", errors="ignore", newline="") as csv_file:
        rows = list(csv.reader(csv_file))

    if not rows:
        return ""

    return _markdown_table(rows[0], rows[1:])


def _excel_to_markdown(path: Path) -> str:
    try:
        sheets = pd.read_excel(path, sheet_name=None, dtype=str, keep_default_na=False)
    except ImportError as exc:
        raise ValueError("Excel 转 Markdown 需要安装 openpyxl/xlrd 依赖") from exc

    markdown_parts = []
    for sheet_name, dataframe in sheets.items():
        markdown_parts.append(f"## {sheet_name}\n\n{_dataframe_to_markdown(dataframe)}")

    return "\n\n".join(markdown_parts)


def _docx_table_to_markdown(table) -> str:
    rows = [
        [cell.text.strip() for cell in row.cells]
        for row in table.rows
    ]
    if not rows:
        return ""

    return _markdown_table(rows[0], rows[1:])


def _docx_to_markdown(path: Path) -> str:
    document = DocxDocument(str(path))
    markdown_parts = []

    for paragraph in document.paragraphs:
        text_value = paragraph.text.strip()
        if text_value:
            markdown_parts.append(text_value)

    for table in document.tables:
        table_markdown = _docx_table_to_markdown(table)
        if table_markdown:
            markdown_parts.append(table_markdown)

    return "\n\n".join(markdown_parts)


def _pdf_table_to_markdown(table: list[list[object]]) -> str:
    rows = table or []
    if not rows:
        return ""

    return _markdown_table(rows[0], rows[1:])


def _pdf_to_markdown(path: Path) -> str:
    markdown_parts = []

    with pdfplumber.open(str(path)) as pdf:
        for page_index, page in enumerate(pdf.pages, 1):
            page_parts = []
            text_value = (page.extract_text() or "").strip()
            if text_value:
                page_parts.append(text_value)

            for table in page.extract_tables() or []:
                table_markdown = _pdf_table_to_markdown(table)
                if table_markdown:
                    page_parts.append(table_markdown)

            if page_parts:
                markdown_parts.append(
                    f"## 第 {page_index} 页\n\n" + "\n\n".join(page_parts)
                )

    return "\n\n".join(markdown_parts)


def _find_office_binary() -> str | None:
    for candidate in (
        "soffice",
        "libreoffice",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ):
        resolved = shutil.which(candidate)
        if resolved:
            return resolved

        candidate_path = Path(candidate)
        if candidate_path.exists():
            return str(candidate_path)

    return None


def _extract_legacy_doc_text(path: Path) -> str:
    data = path.read_bytes()
    text_candidates = []

    for encoding in ("utf-16le", "gb18030", "latin1"):
        decoded = data.decode(encoding, errors="ignore")
        decoded = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]+", "\n", decoded)
        for line in decoded.splitlines():
            cleaned = " ".join(line.split())
            if len(cleaned) >= 12 and re.search(r"[\u4e00-\u9fffA-Za-z0-9]", cleaned):
                text_candidates.append(cleaned)

    unique_lines = []
    seen = set()
    for line in text_candidates:
        if line in seen:
            continue
        seen.add(line)
        unique_lines.append(line)

    return "\n\n".join(unique_lines)


def _doc_to_markdown(path: Path) -> str:
    office_binary = _find_office_binary()
    if office_binary:
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                subprocess.run(
                    [
                        office_binary,
                        "--headless",
                        "--convert-to",
                        "docx",
                        "--outdir",
                        temp_dir,
                        str(path),
                    ],
                    check=True,
                    capture_output=True,
                    timeout=90,
                )
                converted_path = Path(temp_dir) / f"{path.stem}.docx"
                if converted_path.exists():
                    return _docx_to_markdown(converted_path)
        except Exception:
            pass

    return _extract_legacy_doc_text(path)


def convert_file_to_markdown(path: Path, original_name: str) -> str:
    extension = _extension(original_name)

    if extension in {".txt", ".md"}:
        markdown = _text_file_to_markdown(path)
    elif extension == ".csv":
        markdown = _csv_to_markdown(path)
    elif extension in {".xls", ".xlsx"}:
        markdown = _excel_to_markdown(path)
    elif extension == ".docx":
        markdown = _docx_to_markdown(path)
    elif extension == ".doc":
        markdown = _doc_to_markdown(path)
    elif extension == ".pdf":
        markdown = _pdf_to_markdown(path)
    else:
        raise ValueError(f"不支持的文件类型: {extension}")

    markdown = markdown.strip()
    if not markdown:
        raise ValueError("文件未提取到可入库文本")

    return f"# {original_name}\n\n{markdown}\n"


def _row_to_rag_file(row: Mapping[str, object]) -> RagFileRecord:
    return RagFileRecord(
        id=int(row["id"]),
        original_name=str(row["original_name"]),
        original_extension=str(row["original_extension"]),
        mime_type=str(row["mime_type"] or ""),
        size_bytes=int(row["size_bytes"] or 0),
        sha256=str(row["sha256"]),
        source_path=str(row["source_path"]),
        markdown_path=str(row["markdown_path"]),
        markdown_content=str(row.get("markdown_content") or ""),
        status=str(row["status"]),
        error_message=(
            str(row["error_message"])
            if row.get("error_message") is not None
            else None
        ),
        chunk_count=int(row["chunk_count"] or 0),
        created_at=_isoformat(row["created_at"]),
        updated_at=_isoformat(row["updated_at"]),
    )


class RagFileService:
    def __init__(self) -> None:
        self._engine = create_engine(
            get_database_url(),
            pool_pre_ping=True,
            pool_recycle=1800,
            future=True,
        )
        self.data_path = self._resolve_data_path()
        self.original_dir = self.data_path / "uploaded"
        self.markdown_dir = self.data_path / "markdown"
        self.temp_dir = self.data_path / "tmp"
        self.original_dir.mkdir(parents=True, exist_ok=True)
        self.markdown_dir.mkdir(parents=True, exist_ok=True)
        self.temp_dir.mkdir(parents=True, exist_ok=True)
        self._initialize_database()

    @staticmethod
    def allowed_file_types() -> list[str]:
        return [extension.lstrip(".") for extension in ALLOWED_RAG_UPLOAD_EXTENSIONS]

    @staticmethod
    def _resolve_data_path() -> Path:
        path_value = str(chroma_conf["data_path"])
        path = Path(path_value)
        if path.is_absolute():
            return path

        return (PROJECT_ROOT / path).resolve()

    def _initialize_database(self) -> None:
        with self._engine.begin() as connection:
            connection.execute(text(CREATE_RAG_FILES_SQL))

    def create_temp_upload_path(self, filename: str) -> Path:
        safe_name = _safe_original_name(filename)
        extension = _extension(safe_name)
        self._validate_extension(extension)
        return self.temp_dir / f"{datetime.now(timezone.utc).timestamp():.6f}-{safe_name}"

    def prepare_uploaded_file(
        self,
        original_name: str,
        uploaded_path: Path,
        mime_type: str | None = None,
    ) -> PreparedRagFile:
        safe_name = _safe_original_name(original_name)
        extension = _extension(safe_name)
        self._validate_extension(extension)

        sha256 = get_file_sha256_hex(uploaded_path)
        if not sha256:
            raise ValueError("文件 SHA256 计算失败")

        existing = self.get_file_by_sha256(sha256, include_content=False)
        if existing is not None and existing.status in {
            RAG_FILE_STATUS_CONVERTING,
            RAG_FILE_STATUS_INDEXING,
            RAG_FILE_STATUS_READY,
        }:
            return PreparedRagFile(record=existing, should_process=False)

        source_path = self.original_dir / f"{sha256}{extension}"
        markdown_path = self.markdown_dir / f"{sha256}.md"
        size_bytes = uploaded_path.stat().st_size

        if not source_path.exists():
            shutil.copyfile(uploaded_path, source_path)

        now = _now_utc()
        with self._engine.begin() as connection:
            row = connection.execute(
                text("SELECT id FROM rag_files WHERE sha256 = :sha256 LIMIT 1"),
                {"sha256": sha256},
            ).mappings().fetchone()

            payload = {
                "original_name": safe_name,
                "original_extension": extension.lstrip("."),
                "mime_type": (mime_type or "")[:255],
                "size_bytes": size_bytes,
                "sha256": sha256,
                "source_path": _relative_project_path(source_path),
                "markdown_path": _relative_project_path(markdown_path),
                "markdown_content": "",
                "status": RAG_FILE_STATUS_PENDING,
                "error_message": None,
                "now": now,
            }

            if row is None:
                result = connection.execute(
                    text(
                        """
                        INSERT INTO rag_files (
                            original_name,
                            original_extension,
                            mime_type,
                            size_bytes,
                            sha256,
                            source_path,
                            markdown_path,
                            markdown_content,
                            status,
                            error_message,
                            created_at,
                            updated_at
                        )
                        VALUES (
                            :original_name,
                            :original_extension,
                            :mime_type,
                            :size_bytes,
                            :sha256,
                            :source_path,
                            :markdown_path,
                            :markdown_content,
                            :status,
                            :error_message,
                            :now,
                            :now
                        )
                        """
                    ),
                    payload,
                )
                file_id = int(result.lastrowid)
            else:
                file_id = int(row["id"])
                connection.execute(
                    text(
                        """
                        UPDATE rag_files
                        SET
                            original_name = :original_name,
                            original_extension = :original_extension,
                            mime_type = :mime_type,
                            size_bytes = :size_bytes,
                            source_path = :source_path,
                            markdown_path = :markdown_path,
                            markdown_content = :markdown_content,
                            status = :status,
                            error_message = :error_message,
                            chunk_count = 0,
                            updated_at = :now
                        WHERE id = :file_id
                        """
                    ),
                    {**payload, "file_id": file_id},
                )

        record = self.get_file(file_id, include_content=False)
        if record is None:
            raise FileNotFoundError("文件记录创建失败")

        return PreparedRagFile(record=record, should_process=True)

    def convert_file_to_markdown_record(self, file_id: int) -> RagFileRecord:
        record = self.get_file(file_id, include_content=False)
        if record is None:
            raise FileNotFoundError("文件不存在")

        self.update_file_status(
            file_id,
            RAG_FILE_STATUS_CONVERTING,
            error_message=None,
            chunk_count=0,
        )

        source_path = _resolve_project_path(record.source_path)
        markdown_path = _resolve_project_path(record.markdown_path)
        markdown_path.parent.mkdir(parents=True, exist_ok=True)

        try:
            markdown_content = convert_file_to_markdown(source_path, record.original_name)
            markdown_path.write_text(markdown_content, encoding="utf-8")
        except Exception as exc:
            error_message = str(exc)[:2000]
            markdown_path.write_text("", encoding="utf-8")
            self.mark_file_failed(file_id, error_message)
            raise

        now = _now_utc()
        with self._engine.begin() as connection:
            connection.execute(
                text(
                    """
                    UPDATE rag_files
                    SET
                        markdown_content = :markdown_content,
                        status = :status,
                        error_message = NULL,
                        chunk_count = 0,
                        updated_at = :now
                    WHERE id = :file_id
                    """
                ),
                {
                    "file_id": file_id,
                    "markdown_content": markdown_content,
                    "status": RAG_FILE_STATUS_INDEXING,
                    "now": now,
                },
            )

        converted = self.get_file(file_id, include_content=True)
        if converted is None:
            raise FileNotFoundError("文件不存在")

        return converted

    def list_files(self, include_content: bool = False) -> list[RagFileRecord]:
        markdown_select = "markdown_content" if include_content else "'' AS markdown_content"
        with self._engine.connect() as connection:
            rows = connection.execute(
                text(
                    f"""
                    SELECT
                        id,
                        original_name,
                        original_extension,
                        mime_type,
                        size_bytes,
                        sha256,
                        source_path,
                        markdown_path,
                        {markdown_select},
                        status,
                        error_message,
                        chunk_count,
                        created_at,
                        updated_at
                    FROM rag_files
                    ORDER BY updated_at DESC, id DESC
                    """
                )
            ).mappings().fetchall()

        return [_row_to_rag_file(row) for row in rows]

    def list_ready_files(self) -> list[RagFileRecord]:
        with self._engine.connect() as connection:
            rows = connection.execute(
                text(
                    """
                    SELECT
                        id,
                        original_name,
                        original_extension,
                        mime_type,
                        size_bytes,
                        sha256,
                        source_path,
                        markdown_path,
                        markdown_content,
                        status,
                        error_message,
                        chunk_count,
                        created_at,
                        updated_at
                    FROM rag_files
                    WHERE status = 'ready'
                    ORDER BY updated_at DESC, id DESC
                    """
                )
            ).mappings().fetchall()

        return [_row_to_rag_file(row) for row in rows]

    def list_indexable_files(self) -> list[RagFileRecord]:
        with self._engine.connect() as connection:
            rows = connection.execute(
                text(
                    """
                    SELECT
                        id,
                        original_name,
                        original_extension,
                        mime_type,
                        size_bytes,
                        sha256,
                        source_path,
                        markdown_path,
                        markdown_content,
                        status,
                        error_message,
                        chunk_count,
                        created_at,
                        updated_at
                    FROM rag_files
                    WHERE status IN ('ready', 'indexing')
                    ORDER BY updated_at DESC, id DESC
                    """
                )
            ).mappings().fetchall()

        return [_row_to_rag_file(row) for row in rows]

    def get_file(
        self,
        file_id: int,
        include_content: bool = True,
    ) -> RagFileRecord | None:
        markdown_select = "markdown_content" if include_content else "'' AS markdown_content"
        with self._engine.connect() as connection:
            row = connection.execute(
                text(
                    f"""
                    SELECT
                        id,
                        original_name,
                        original_extension,
                        mime_type,
                        size_bytes,
                        sha256,
                        source_path,
                        markdown_path,
                        {markdown_select},
                        status,
                        error_message,
                        chunk_count,
                        created_at,
                        updated_at
                    FROM rag_files
                    WHERE id = :file_id
                    LIMIT 1
                    """
                ),
                {"file_id": file_id},
            ).mappings().fetchone()

        return _row_to_rag_file(row) if row is not None else None

    def get_file_by_sha256(
        self,
        sha256: str,
        include_content: bool = True,
    ) -> RagFileRecord | None:
        markdown_select = "markdown_content" if include_content else "'' AS markdown_content"
        with self._engine.connect() as connection:
            row = connection.execute(
                text(
                    f"""
                    SELECT
                        id,
                        original_name,
                        original_extension,
                        mime_type,
                        size_bytes,
                        sha256,
                        source_path,
                        markdown_path,
                        {markdown_select},
                        status,
                        error_message,
                        chunk_count,
                        created_at,
                        updated_at
                    FROM rag_files
                    WHERE sha256 = :sha256
                    LIMIT 1
                    """
                ),
                {"sha256": sha256},
            ).mappings().fetchone()

        return _row_to_rag_file(row) if row is not None else None

    def update_chunk_count(self, file_id: int, chunk_count: int) -> None:
        with self._engine.begin() as connection:
            connection.execute(
                text(
                    """
                    UPDATE rag_files
                    SET chunk_count = :chunk_count
                    WHERE id = :file_id
                    """
                ),
                {
                    "file_id": file_id,
                    "chunk_count": max(0, int(chunk_count)),
                },
            )

    def update_file_status(
        self,
        file_id: int,
        status_value: str,
        error_message: str | None = None,
        chunk_count: int | None = None,
    ) -> None:
        assignments = [
            "status = :status",
            "error_message = :error_message",
            "updated_at = :now",
        ]
        payload: dict[str, object] = {
            "file_id": int(file_id),
            "status": status_value,
            "error_message": error_message[:2000] if error_message else None,
            "now": _now_utc(),
        }

        if chunk_count is not None:
            assignments.append("chunk_count = :chunk_count")
            payload["chunk_count"] = max(0, int(chunk_count))

        with self._engine.begin() as connection:
            connection.execute(
                text(
                    f"""
                    UPDATE rag_files
                    SET {", ".join(assignments)}
                    WHERE id = :file_id
                    """
                ),
                payload,
            )

    def mark_file_ready(self, file_id: int) -> None:
        self.update_file_status(
            file_id,
            RAG_FILE_STATUS_READY,
            error_message=None,
        )

    def mark_file_failed(self, file_id: int, error_message: str) -> None:
        self.update_file_status(
            file_id,
            RAG_FILE_STATUS_FAILED,
            error_message=error_message,
            chunk_count=0,
        )

    def mark_files_indexing(self, file_ids: Iterable[int]) -> None:
        self._update_files_status(
            file_ids,
            RAG_FILE_STATUS_INDEXING,
            error_message=None,
        )

    def mark_files_ready(self, file_ids: Iterable[int]) -> None:
        self._update_files_status(
            file_ids,
            RAG_FILE_STATUS_READY,
            error_message=None,
        )

    def _update_files_status(
        self,
        file_ids: Iterable[int],
        status_value: str,
        error_message: str | None = None,
    ) -> None:
        normalized_ids = sorted({int(file_id) for file_id in file_ids})
        if not normalized_ids:
            return

        placeholders = ", ".join(
            f":file_id_{index}" for index, _ in enumerate(normalized_ids)
        )
        payload: dict[str, object] = {
            "status": status_value,
            "error_message": error_message[:2000] if error_message else None,
            "now": _now_utc(),
        }
        payload.update(
            {
                f"file_id_{index}": file_id
                for index, file_id in enumerate(normalized_ids)
            }
        )

        with self._engine.begin() as connection:
            connection.execute(
                text(
                    f"""
                    UPDATE rag_files
                    SET
                        status = :status,
                        error_message = :error_message,
                        updated_at = :now
                    WHERE id IN ({placeholders})
                    """
                ),
                payload,
            )

    def delete_file(self, file_id: int) -> None:
        record = self.get_file(file_id, include_content=False)
        if record is None:
            raise FileNotFoundError("文件不存在")

        with self._engine.begin() as connection:
            connection.execute(
                text("DELETE FROM rag_files WHERE id = :file_id"),
                {"file_id": file_id},
            )

        for path_value in (record.source_path, record.markdown_path):
            try:
                path = _resolve_project_path(path_value)
            except ValueError:
                continue

            if path.exists() and path.is_file():
                path.unlink()

    @staticmethod
    def _validate_extension(extension: str) -> None:
        if extension not in ALLOWED_RAG_UPLOAD_EXTENSIONS:
            allowed = ", ".join(extension.lstrip(".") for extension in ALLOWED_RAG_UPLOAD_EXTENSIONS)
            raise ValueError(f"仅支持上传以下文件类型: {allowed}")


@lru_cache(maxsize=1)
def get_rag_file_service() -> RagFileService:
    return RagFileService()
