import csv
import json
import os
import re
import shutil
import subprocess
import tempfile
from collections.abc import Callable, Iterable, Mapping
from dataclasses import dataclass
from datetime import datetime, timezone
from functools import lru_cache
from pathlib import Path
from typing import Any

import pandas as pd
import pdfplumber
from docx import Document as DocxDocument
from PIL import Image, ImageOps
from pypdf import PdfReader, PdfWriter
from sqlalchemy import text

from api.core.settings import PROJECT_ROOT
from db.core import get_database_engine
from db.schema import initialize_rag_files_schema
from rag.paddle_ocr_service import (
    PaddleOcrClient,
    get_paddle_ocr_config_service,
)
from utils.config_handler import chroma_conf
from utils.file_handler import get_file_sha256_hex


ALLOWED_RAG_UPLOAD_EXTENSIONS = (
    ".doc",
    ".docx",
    ".txt",
    ".pdf",
    ".csv",
    ".xls",
    ".xlsx",
    ".png",
    ".jpg",
    ".jpeg",
)

RAG_FILE_STATUS_PENDING = "pending"
RAG_FILE_STATUS_EXTRACTING = "extracting"
RAG_FILE_STATUS_OCR = "ocr"
RAG_FILE_STATUS_RENDERING = "rendering"
RAG_FILE_STATUS_INDEXING = "indexing"
RAG_FILE_STATUS_READY = "ready"
RAG_FILE_STATUS_FAILED = "failed"
RAG_FILE_INDEXABLE_STATUSES = (
    RAG_FILE_STATUS_READY,
    RAG_FILE_STATUS_INDEXING,
)

STRUCTURED_CONTENT_VERSION = 1
TABLE_ROWS_PER_BLOCK = 24
PDF_PAGE_OCR_CHARACTER_THRESHOLD = 40


@dataclass(frozen=True)
class RagFileRecord:
    id: int
    original_name: str
    original_extension: str
    mime_type: str
    size_bytes: int
    sha256: str
    source_path: str
    content_json_path: str
    preview_pdf_path: str
    content_json: str
    status: str
    error_message: str | None
    chunk_count: int
    used_ocr: bool
    created_at: str
    updated_at: str


@dataclass(frozen=True)
class PreparedRagFile:
    record: RagFileRecord
    should_process: bool


@dataclass(frozen=True)
class PdfPageExtraction:
    page_number: int
    blocks: list[dict[str, Any]]
    needs_ocr: bool


def _now_utc() -> datetime:
    return datetime.now(timezone.utc).replace(tzinfo=None)


def _isoformat(value: object) -> str:
    if isinstance(value, datetime):
        normalized = value
    else:
        normalized = _now_utc()

    if normalized.tzinfo is None:
        normalized = normalized.replace(tzinfo=timezone.utc)

    return normalized.astimezone(timezone.utc).isoformat().replace("+00:00", "Z")


def _safe_original_name(filename: str) -> str:
    safe_name = Path(filename).name.strip()
    if not safe_name or safe_name in {".", ".."}:
        raise ValueError("文件名无效")

    if len(safe_name) > 255:
        raise ValueError("文件名不能超过 255 个字符")

    return safe_name


def _extension(filename: str) -> str:
    return Path(filename).suffix.lower()


def _relative_project_path(path: Path) -> str:
    return path.resolve().relative_to(PROJECT_ROOT.resolve()).as_posix()


def _resolve_project_path(path_value: str) -> Path:
    path = Path(path_value)
    if path.is_absolute():
        target = path.resolve()
    else:
        target = (PROJECT_ROOT / path).resolve()

    target.relative_to(PROJECT_ROOT.resolve())
    return target


def _clean_text(value: object) -> str:
    return str(value or "").replace("\r\n", "\n").replace("\r", "\n").strip()


def _clean_cell(value: object) -> str:
    return " ".join(str(value or "").replace("\r", " ").replace("\n", " ").split())


def _normalize_table_rows(rows: Iterable[Iterable[object]]) -> list[list[str]]:
    cleaned_rows = [[_clean_cell(cell) for cell in row] for row in rows]
    cleaned_rows = [row for row in cleaned_rows if any(row)]
    if not cleaned_rows:
        return []

    max_columns = max(len(row) for row in cleaned_rows)
    padded_rows = [row + [""] * (max_columns - len(row)) for row in cleaned_rows]
    keep_columns = [
        column_index
        for column_index in range(max_columns)
        if any(row[column_index] for row in padded_rows)
    ]
    if not keep_columns:
        return []

    return [[row[column_index] for column_index in keep_columns] for row in padded_rows]


def _table_to_text(rows: Iterable[Iterable[object]]) -> str:
    lines = []
    for row in _normalize_table_rows(rows):
        cells = [cell for cell in row if cell]
        if cells:
            lines.append(" | ".join(cells))

    return "\n".join(lines).strip()


def _make_block(
    page_number: int,
    block_index: int,
    block_type: str,
    text_value: str,
    bbox: list[float] | None = None,
    extra: Mapping[str, Any] | None = None,
) -> dict[str, Any] | None:
    normalized_text = _clean_text(text_value)
    if not normalized_text:
        return None

    block = {
        "id": f"p{page_number}-b{block_index}",
        "type": block_type,
        "text": normalized_text,
        "page_number": max(1, int(page_number)),
        "bbox": bbox,
    }
    if extra:
        block.update(dict(extra))

    return block


def _make_table_blocks(
    page_number: int,
    start_block_index: int,
    rows: Iterable[Iterable[object]],
    title: str | None = None,
) -> list[dict[str, Any]]:
    normalized_rows = _normalize_table_rows(rows)
    if not normalized_rows:
        return []

    header = normalized_rows[0]
    body_rows = normalized_rows[1:]
    row_groups = [normalized_rows] if not body_rows else [
        [header, *body_rows[index:index + TABLE_ROWS_PER_BLOCK]]
        for index in range(0, len(body_rows), TABLE_ROWS_PER_BLOCK)
    ]

    blocks = []
    for group_index, row_group in enumerate(row_groups):
        block_title = ""
        if title:
            block_title = title if group_index == 0 else f"{title}（续 {group_index + 1}）"

        table_text = _table_to_text(row_group)
        text_value = f"{block_title}\n{table_text}" if block_title else table_text
        block = _make_block(
            page_number,
            start_block_index + group_index,
            "table",
            text_value,
            extra={
                "title": block_title,
                "rows": row_group,
            },
        )
        if block:
            blocks.append(block)

    return blocks


def _new_structured_content(
    original_name: str,
    extension: str,
    sha256: str,
    pages: list[dict[str, Any]],
) -> dict[str, Any]:
    return {
        "version": STRUCTURED_CONTENT_VERSION,
        "file": {
            "name": original_name,
            "extension": extension.lstrip("."),
            "sha256": sha256,
        },
        "pages": pages,
    }


def _iter_blocks(content: Mapping[str, Any]) -> Iterable[dict[str, Any]]:
    pages = content.get("pages")
    if not isinstance(pages, list):
        return

    for page in pages:
        if not isinstance(page, dict):
            continue

        blocks = page.get("blocks")
        if not isinstance(blocks, list):
            continue

        for block in blocks:
            if isinstance(block, dict):
                yield block


def structured_content_text_blocks(content_json: str) -> list[dict[str, Any]]:
    try:
        content = json.loads(content_json or "{}")
    except json.JSONDecodeError:
        return []

    blocks = []
    for block in _iter_blocks(content):
        text_value = _clean_text(block.get("text"))
        if not text_value:
            continue

        page_number = block.get("page_number")
        try:
            normalized_page_number = int(page_number)
        except (TypeError, ValueError):
            normalized_page_number = 1

        bbox = block.get("bbox")
        blocks.append(
            {
                "id": str(block.get("id") or ""),
                "type": str(block.get("type") or "paragraph"),
                "text": text_value,
                "page_number": max(1, normalized_page_number),
                "bbox": bbox if isinstance(bbox, list) else None,
            }
        )

    return blocks


def _paragraph_blocks_from_text(
    text_value: str,
    page_number: int = 1,
    block_type: str = "paragraph",
) -> list[dict[str, Any]]:
    chunks = [
        chunk.strip()
        for chunk in re.split(r"\n\s*\n", _clean_text(text_value))
        if chunk.strip()
    ]

    if not chunks:
        chunks = [line.strip() for line in _clean_text(text_value).splitlines() if line.strip()]

    blocks = []
    for index, chunk in enumerate(chunks, 1):
        block = _make_block(page_number, index, block_type, chunk)
        if block:
            blocks.append(block)

    return blocks


def _read_text(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue

    return path.read_text(encoding="utf-8", errors="ignore")


def _text_to_structured_content(
    path: Path,
    original_name: str,
    extension: str,
    sha256: str,
) -> dict[str, Any]:
    blocks = _paragraph_blocks_from_text(_read_text(path), page_number=1)
    return _new_structured_content(
        original_name,
        extension,
        sha256,
        [{"page_number": 1, "blocks": blocks}],
    )


def _csv_to_structured_content(
    path: Path,
    original_name: str,
    extension: str,
    sha256: str,
) -> dict[str, Any]:
    rows: list[list[object]]
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            dataframe = pd.read_csv(
                path,
                dtype=str,
                keep_default_na=False,
                encoding=encoding,
                header=None,
            )
            rows = dataframe.fillna("").values.tolist()
            break
        except UnicodeDecodeError:
            continue
        except Exception:
            rows = []
            break
    else:
        rows = []

    if not rows:
        with path.open("r", encoding="utf-8", errors="ignore", newline="") as csv_file:
            rows = list(csv.reader(csv_file))

    blocks = _make_table_blocks(1, 1, rows)
    return _new_structured_content(
        original_name,
        extension,
        sha256,
        [{"page_number": 1, "blocks": blocks}],
    )


def _excel_to_structured_content(
    path: Path,
    original_name: str,
    extension: str,
    sha256: str,
) -> dict[str, Any]:
    try:
        sheets = pd.read_excel(
            path,
            sheet_name=None,
            dtype=str,
            keep_default_na=False,
            header=None,
        )
    except ImportError as exc:
        raise ValueError("Excel 解析需要安装 openpyxl/xlrd 依赖") from exc

    pages = []
    for page_index, (sheet_name, dataframe) in enumerate(sheets.items(), 1):
        rows = dataframe.fillna("").values.tolist()
        blocks = _make_table_blocks(page_index, 1, rows, title=f"工作表：{sheet_name}")
        pages.append(
            {
                "page_number": page_index,
                "blocks": blocks,
            }
        )

    return _new_structured_content(original_name, extension, sha256, pages)


def _docx_to_blocks(path: Path) -> list[dict[str, Any]]:
    document = DocxDocument(str(path))
    blocks = []
    block_index = 1

    for paragraph in document.paragraphs:
        block = _make_block(1, block_index, "paragraph", paragraph.text)
        if block:
            blocks.append(block)
            block_index += 1

    for table in document.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        table_blocks = _make_table_blocks(1, block_index, rows)
        blocks.extend(table_blocks)
        block_index += len(table_blocks)

    return blocks


def _docx_to_structured_content(
    path: Path,
    original_name: str,
    extension: str,
    sha256: str,
) -> dict[str, Any]:
    return _new_structured_content(
        original_name,
        extension,
        sha256,
        [{"page_number": 1, "blocks": _docx_to_blocks(path)}],
    )


def _find_office_binary() -> str | None:
    for candidate in (
        "soffice",
        "libreoffice",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ):
        resolved = shutil.which(candidate)
        if resolved:
            return resolved

        candidate_path = Path(candidate)
        if candidate_path.exists():
            return str(candidate_path)

    return None


def _extract_legacy_doc_text(path: Path) -> str:
    data = path.read_bytes()
    text_candidates = []

    for encoding in ("utf-16le", "gb18030", "latin1"):
        decoded = data.decode(encoding, errors="ignore")
        decoded = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]+", "\n", decoded)
        for line in decoded.splitlines():
            cleaned = " ".join(line.split())
            if len(cleaned) >= 12 and re.search(r"[\u4e00-\u9fffA-Za-z0-9]", cleaned):
                text_candidates.append(cleaned)

    unique_lines = []
    seen = set()
    for line in text_candidates:
        if line in seen:
            continue
        seen.add(line)
        unique_lines.append(line)

    return "\n\n".join(unique_lines)


def _doc_to_structured_content(
    path: Path,
    original_name: str,
    extension: str,
    sha256: str,
) -> dict[str, Any]:
    office_binary = _find_office_binary()
    if office_binary:
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                subprocess.run(
                    [
                        office_binary,
                        "--headless",
                        "--convert-to",
                        "docx",
                        "--outdir",
                        temp_dir,
                        str(path),
                    ],
                    check=True,
                    capture_output=True,
                    timeout=90,
                )
                converted_path = Path(temp_dir) / f"{path.stem}.docx"
                if converted_path.exists():
                    return _docx_to_structured_content(
                        converted_path,
                        original_name,
                        extension,
                        sha256,
                    )
        except Exception:
            pass

    blocks = _paragraph_blocks_from_text(_extract_legacy_doc_text(path), page_number=1)
    return _new_structured_content(
        original_name,
        extension,
        sha256,
        [{"page_number": 1, "blocks": blocks}],
    )


def _block_text_character_count(blocks: Iterable[Mapping[str, Any]]) -> int:
    return sum(
        len(re.sub(r"\s+", "", str(block.get("text") or "")))
        for block in blocks
    )


def _pdf_page_visual_object_count(page: Any) -> int:
    return sum(
        len(getattr(page, attribute_name, None) or [])
        for attribute_name in ("images", "curves", "rects", "lines")
    )


def _extract_pdf_page_blocks(page: Any, page_number: int) -> list[dict[str, Any]]:
    blocks = []
    block_index = 1
    text_value = (page.extract_text() or "").strip()
    if text_value:
        block = _make_block(page_number, block_index, "page_text", text_value)
        if block:
            blocks.append(block)
            block_index += 1

    for table in page.extract_tables() or []:
        table_blocks = _make_table_blocks(page_number, block_index, table)
        blocks.extend(table_blocks)
        block_index += len(table_blocks)

    return blocks


def _pdf_page_needs_ocr(page: Any, blocks: list[dict[str, Any]]) -> bool:
    extracted_character_count = _block_text_character_count(blocks)
    if extracted_character_count >= PDF_PAGE_OCR_CHARACTER_THRESHOLD:
        return False

    visual_object_count = _pdf_page_visual_object_count(page)
    if visual_object_count <= 0:
        return False

    has_text_layer = bool(getattr(page, "chars", None) or [])
    has_raster_image = bool(getattr(page, "images", None) or [])
    if extracted_character_count == 0:
        return True

    return has_raster_image or not has_text_layer


def _extract_pdf_pages(path: Path) -> list[PdfPageExtraction]:
    page_extractions = []
    with pdfplumber.open(str(path)) as pdf:
        for page_number, page in enumerate(pdf.pages, 1):
            blocks = _extract_pdf_page_blocks(page, page_number)
            page_extractions.append(
                PdfPageExtraction(
                    page_number=page_number,
                    blocks=blocks,
                    needs_ocr=_pdf_page_needs_ocr(page, blocks),
                )
            )

    return page_extractions


def _write_pdf_page_subset(
    source_path: Path,
    page_numbers: list[int],
    output_path: Path,
) -> None:
    reader = PdfReader(str(source_path))
    writer = PdfWriter()
    for page_number in page_numbers:
        writer.add_page(reader.pages[page_number - 1])

    with output_path.open("wb") as output_file:
        writer.write(output_file)


def extract_paddle_ocr_pdf_pages(
    path: Path,
    page_numbers: list[int],
) -> tuple[str, dict[int, str]]:
    config = get_paddle_ocr_config_service().get_config()
    client = PaddleOcrClient(config.api_key, model_name=config.model_name)

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir) / "ocr-pages.pdf"
        _write_pdf_page_subset(path, page_numbers, temp_path)
        job_id, page_texts = client.extract_pages(temp_path)

    return job_id, {
        page_number: _clean_text(page_texts[index])
        if index < len(page_texts)
        else ""
        for index, page_number in enumerate(page_numbers)
    }


def _pdf_to_structured_content(
    path: Path,
    original_name: str,
    extension: str,
    sha256: str,
    ocr_status_callback: Callable[[], None] | None = None,
) -> dict[str, Any]:
    page_extractions = _extract_pdf_pages(path)
    ocr_page_numbers = [
        page_extraction.page_number
        for page_extraction in page_extractions
        if page_extraction.needs_ocr
    ]
    ocr_text_by_page: dict[int, str] = {}

    if ocr_page_numbers:
        if ocr_status_callback:
            ocr_status_callback()
        _, ocr_text_by_page = extract_paddle_ocr_pdf_pages(path, ocr_page_numbers)

    pages = []
    for page_extraction in page_extractions:
        blocks = page_extraction.blocks
        ocr_text = ocr_text_by_page.get(page_extraction.page_number)
        if ocr_text:
            block = _make_block(
                page_extraction.page_number,
                1,
                "ocr_markdown",
                ocr_text,
            )
            blocks = [block] if block else []

        pages.append(
            {
                "page_number": page_extraction.page_number,
                "blocks": blocks,
            }
        )

    content = _new_structured_content(original_name, extension, sha256, pages)
    if ocr_page_numbers:
        content["ocr"] = {
            "page_numbers": ocr_page_numbers,
        }

    return content


def _pdf_has_ocr_pages(path: Path) -> bool:
    return any(page_extraction.needs_ocr for page_extraction in _extract_pdf_pages(path))


def requires_paddle_ocr(path: Path, original_name: str) -> bool:
    extension = _extension(original_name)
    if extension in {".png", ".jpg", ".jpeg"}:
        return True
    if extension == ".pdf":
        return _pdf_has_ocr_pages(path)
    return False


def _ocr_pages_to_structured_content(
    page_texts: list[str],
    original_name: str,
    extension: str,
    sha256: str,
) -> dict[str, Any]:
    pages = []
    for page_number, text_value in enumerate(page_texts, 1):
        block = _make_block(
            page_number,
            1,
            "ocr_markdown",
            text_value,
            bbox=None,
        )
        pages.append(
            {
                "page_number": page_number,
                "blocks": [block] if block else [],
            }
        )

    content = _new_structured_content(original_name, extension, sha256, pages)
    if not structured_content_text_blocks(
        json.dumps(content, ensure_ascii=False, separators=(",", ":"))
    ):
        raise ValueError("PaddleOCR 未识别到可用于 RAG 的文字")

    return content


def extract_paddle_ocr_content(
    path: Path,
    original_name: str,
    sha256: str,
) -> tuple[str, dict[str, Any]]:
    config = get_paddle_ocr_config_service().get_config()
    client = PaddleOcrClient(config.api_key, model_name=config.model_name)
    job_id, page_texts = client.extract_pages(path)
    return job_id, _ocr_pages_to_structured_content(
        page_texts,
        original_name,
        _extension(original_name),
        sha256,
    )


def extract_structured_content(
    path: Path,
    original_name: str,
    sha256: str,
    ocr_status_callback: Callable[[], None] | None = None,
) -> dict[str, Any]:
    extension = _extension(original_name)

    if extension == ".txt":
        content = _text_to_structured_content(path, original_name, extension, sha256)
    elif extension == ".csv":
        content = _csv_to_structured_content(path, original_name, extension, sha256)
    elif extension in {".xls", ".xlsx"}:
        content = _excel_to_structured_content(path, original_name, extension, sha256)
    elif extension == ".docx":
        content = _docx_to_structured_content(path, original_name, extension, sha256)
    elif extension == ".doc":
        content = _doc_to_structured_content(path, original_name, extension, sha256)
    elif extension == ".pdf":
        content = _pdf_to_structured_content(
            path,
            original_name,
            extension,
            sha256,
            ocr_status_callback=ocr_status_callback,
        )
    else:
        raise ValueError(f"不支持的文件类型: {extension}")

    if not structured_content_text_blocks(
        json.dumps(content, ensure_ascii=False, separators=(",", ":"))
    ):
        raise ValueError("文件未提取到可入库文本")

    return content


def _pdf_page_count(path: Path) -> int:
    try:
        with pdfplumber.open(str(path)) as pdf:
            return len(pdf.pages)
    except Exception:
        return 0


def _update_preview_metadata(content: dict[str, Any], preview_path: Path) -> dict[str, Any]:
    content["preview"] = {
        "page_count": _pdf_page_count(preview_path),
        "format": "pdf",
    }
    return content


def _copy_pdf_preview(source_path: Path, output_path: Path, content: dict[str, Any]) -> dict[str, Any]:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    if source_path.resolve() != output_path.resolve():
        shutil.copyfile(source_path, output_path)

    return _update_preview_metadata(content, output_path)


def _image_to_pdf_preview(
    source_path: Path,
    output_path: Path,
    content: dict[str, Any],
) -> dict[str, Any]:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with Image.open(source_path) as source_image:
        oriented_image = ImageOps.exif_transpose(source_image)
        if oriented_image.mode in {"RGBA", "LA"} or "transparency" in oriented_image.info:
            rgba_image = oriented_image.convert("RGBA")
            rgb_image = Image.new("RGB", rgba_image.size, "white")
            rgb_image.paste(rgba_image, mask=rgba_image.getchannel("A"))
        else:
            rgb_image = oriented_image.convert("RGB")

        try:
            resolution = float(source_image.info.get("dpi", (96, 96))[0] or 96)
        except (TypeError, ValueError, IndexError):
            resolution = 96
        rgb_image.save(output_path, "PDF", resolution=max(36, resolution))

    return _update_preview_metadata(content, output_path)


def _convert_with_libreoffice_to_pdf(source_path: Path, output_path: Path) -> None:
    office_binary = _find_office_binary()
    if not office_binary:
        raise ValueError("未找到 LibreOffice/soffice，无法生成预览 PDF")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_root = Path(temp_dir)
        temp_input = temp_root / source_path.name
        shutil.copyfile(source_path, temp_input)

        profile_dir = temp_root / "lo-profile"
        profile_dir.mkdir(parents=True, exist_ok=True)
        result = subprocess.run(
            [
                office_binary,
                "--headless",
                "--nologo",
                "--nofirststartwizard",
                f"-env:UserInstallation={profile_dir.as_uri()}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(temp_root),
                str(temp_input),
            ],
            check=False,
            capture_output=True,
            env={**os.environ, "HOME": str(temp_root)},
            text=True,
            timeout=180,
        )
        if result.returncode != 0:
            detail = (result.stderr or result.stdout or "").strip()
            raise ValueError(f"LibreOffice 生成预览 PDF 失败: {detail or result.returncode}")

        converted_path = temp_input.with_suffix(".pdf")
        if not converted_path.exists():
            candidates = sorted(temp_root.glob("*.pdf"))
            converted_path = candidates[0] if candidates else converted_path

        if not converted_path.exists():
            raise ValueError("LibreOffice 未生成预览 PDF")

        shutil.copyfile(converted_path, output_path)


def generate_preview_pdf(
    source_path: Path,
    output_path: Path,
    content: dict[str, Any],
) -> dict[str, Any]:
    extension = _extension(source_path.name)
    if extension == ".pdf":
        return _copy_pdf_preview(source_path, output_path, content)
    if extension in {".png", ".jpg", ".jpeg"}:
        return _image_to_pdf_preview(source_path, output_path, content)

    _convert_with_libreoffice_to_pdf(source_path, output_path)
    return _update_preview_metadata(content, output_path)


def _row_to_rag_file(row: Mapping[str, object]) -> RagFileRecord:
    return RagFileRecord(
        id=int(row["id"]),
        original_name=str(row["original_name"]),
        original_extension=str(row["original_extension"]),
        mime_type=str(row["mime_type"] or ""),
        size_bytes=int(row["size_bytes"] or 0),
        sha256=str(row["sha256"]),
        source_path=str(row["source_path"]),
        content_json_path=str(row["content_json_path"]),
        preview_pdf_path=str(row["preview_pdf_path"]),
        content_json=str(row.get("content_json") or ""),
        status=str(row["status"]),
        error_message=(
            str(row["error_message"])
            if row.get("error_message") is not None
            else None
        ),
        chunk_count=int(row["chunk_count"] or 0),
        used_ocr=bool(row["used_ocr"]),
        created_at=_isoformat(row["created_at"]),
        updated_at=_isoformat(row["updated_at"]),
    )


class RagFileService:
    def __init__(self) -> None:
        self._engine = get_database_engine()
        self.data_path = self._resolve_data_path()
        self.original_dir = self.data_path / "uploaded"
        self.content_dir = self.data_path / "content"
        self.preview_dir = self.data_path / "preview"
        self.temp_dir = self.data_path / "tmp"
        self.original_dir.mkdir(parents=True, exist_ok=True)
        self.content_dir.mkdir(parents=True, exist_ok=True)
        self.preview_dir.mkdir(parents=True, exist_ok=True)
        self.temp_dir.mkdir(parents=True, exist_ok=True)
        self._initialize_database()

    @staticmethod
    def allowed_file_types() -> list[str]:
        return [extension.lstrip(".") for extension in ALLOWED_RAG_UPLOAD_EXTENSIONS]

    @staticmethod
    def _resolve_data_path() -> Path:
        path_value = str(chroma_conf["data_path"])
        path = Path(path_value)
        if path.is_absolute():
            return path

        return (PROJECT_ROOT / path).resolve()

    def _initialize_database(self) -> None:
        initialize_rag_files_schema(self._engine)

    def create_temp_upload_path(self, filename: str) -> Path:
        safe_name = _safe_original_name(filename)
        extension = _extension(safe_name)
        self._validate_extension(extension)
        return self.temp_dir / f"{datetime.now(timezone.utc).timestamp():.6f}-{safe_name}"

    def prepare_uploaded_file(
        self,
        original_name: str,
        uploaded_path: Path,
        mime_type: str | None = None,
    ) -> PreparedRagFile:
        safe_name = _safe_original_name(original_name)
        extension = _extension(safe_name)
        self._validate_extension(extension)

        sha256 = get_file_sha256_hex(uploaded_path)
        if not sha256:
            raise ValueError("文件 SHA256 计算失败")

        existing = self.get_file_by_sha256(sha256, include_content=False)
        if existing is not None and existing.status in {
            RAG_FILE_STATUS_EXTRACTING,
            RAG_FILE_STATUS_OCR,
            RAG_FILE_STATUS_RENDERING,
            RAG_FILE_STATUS_INDEXING,
            RAG_FILE_STATUS_READY,
        }:
            return PreparedRagFile(record=existing, should_process=False)

        source_path = self.original_dir / f"{sha256}{extension}"
        content_json_path = self.content_dir / f"{sha256}.json"
        preview_pdf_path = self.preview_dir / f"{sha256}.pdf"
        size_bytes = uploaded_path.stat().st_size

        if not source_path.exists():
            shutil.copyfile(uploaded_path, source_path)

        now = _now_utc()
        with self._engine.begin() as connection:
            row = connection.execute(
                text("SELECT id FROM rag_files WHERE sha256 = :sha256 LIMIT 1"),
                {"sha256": sha256},
            ).mappings().fetchone()

            payload = {
                "original_name": safe_name,
                "original_extension": extension.lstrip("."),
                "mime_type": (mime_type or "")[:255],
                "size_bytes": size_bytes,
                "sha256": sha256,
                "source_path": _relative_project_path(source_path),
                "content_json_path": _relative_project_path(content_json_path),
                "preview_pdf_path": _relative_project_path(preview_pdf_path),
                "content_json": "",
                "status": RAG_FILE_STATUS_PENDING,
                "error_message": None,
                "now": now,
            }

            if row is None:
                result = connection.execute(
                    text(
                        """
                        INSERT INTO rag_files (
                            original_name,
                            original_extension,
                            mime_type,
                            size_bytes,
                            sha256,
                            source_path,
                            content_json_path,
                            preview_pdf_path,
                            content_json,
                            status,
                            error_message,
                            created_at,
                            updated_at
                        )
                        VALUES (
                            :original_name,
                            :original_extension,
                            :mime_type,
                            :size_bytes,
                            :sha256,
                            :source_path,
                            :content_json_path,
                            :preview_pdf_path,
                            :content_json,
                            :status,
                            :error_message,
                            :now,
                            :now
                        )
                        """
                    ),
                    payload,
                )
                file_id = int(result.lastrowid)
            else:
                file_id = int(row["id"])
                connection.execute(
                    text(
                        """
                        UPDATE rag_files
                        SET
                            original_name = :original_name,
                            original_extension = :original_extension,
                            mime_type = :mime_type,
                            size_bytes = :size_bytes,
                            source_path = :source_path,
                            content_json_path = :content_json_path,
                            preview_pdf_path = :preview_pdf_path,
                            content_json = :content_json,
                            status = :status,
                            error_message = :error_message,
                            chunk_count = 0,
                            used_ocr = 0,
                            updated_at = :now
                        WHERE id = :file_id
                        """
                    ),
                    {**payload, "file_id": file_id},
                )

        record = self.get_file(file_id, include_content=False)
        if record is None:
            raise FileNotFoundError("文件记录创建失败")

        return PreparedRagFile(record=record, should_process=True)

    def process_file_record(self, file_id: int) -> RagFileRecord:
        record = self.get_file(file_id, include_content=False)
        if record is None:
            raise FileNotFoundError("文件不存在")

        self.update_file_status(
            file_id,
            RAG_FILE_STATUS_EXTRACTING,
            error_message=None,
            chunk_count=0,
            used_ocr=False,
        )

        source_path = _resolve_project_path(record.source_path)
        content_json_path = _resolve_project_path(record.content_json_path)
        preview_pdf_path = _resolve_project_path(record.preview_pdf_path)

        try:
            ocr_status_marked = False

            def mark_ocr_status_once() -> None:
                nonlocal ocr_status_marked
                if ocr_status_marked:
                    return

                self.update_file_status(
                    file_id,
                    RAG_FILE_STATUS_OCR,
                    error_message=None,
                    used_ocr=True,
                )
                ocr_status_marked = True

            extension = _extension(record.original_name)
            if extension in {".png", ".jpg", ".jpeg"}:
                mark_ocr_status_once()
                _, structured_content = extract_paddle_ocr_content(
                    source_path,
                    record.original_name,
                    record.sha256,
                )
            else:
                structured_content = extract_structured_content(
                    source_path,
                    record.original_name,
                    record.sha256,
                    ocr_status_callback=mark_ocr_status_once,
                )
            self.update_file_status(file_id, RAG_FILE_STATUS_RENDERING)
            structured_content = generate_preview_pdf(
                source_path,
                preview_pdf_path,
                structured_content,
            )

            content_json = json.dumps(
                structured_content,
                ensure_ascii=False,
                separators=(",", ":"),
            )
            content_json_path.parent.mkdir(parents=True, exist_ok=True)
            content_json_path.write_text(content_json, encoding="utf-8")
        except Exception as exc:
            error_message = str(exc)[:2000]
            self.mark_file_failed(file_id, error_message)
            raise

        now = _now_utc()
        with self._engine.begin() as connection:
            connection.execute(
                text(
                    """
                    UPDATE rag_files
                    SET
                        content_json = :content_json,
                        status = :status,
                        error_message = NULL,
                        chunk_count = 0,
                        updated_at = :now
                    WHERE id = :file_id
                    """
                ),
                {
                    "file_id": file_id,
                    "content_json": content_json,
                    "status": RAG_FILE_STATUS_INDEXING,
                    "now": now,
                },
            )

        processed = self.get_file(file_id, include_content=True)
        if processed is None:
            raise FileNotFoundError("文件不存在")

        return processed

    def list_files(self, include_content: bool = False) -> list[RagFileRecord]:
        content_select = "content_json" if include_content else "'' AS content_json"
        with self._engine.connect() as connection:
            rows = connection.execute(
                text(
                    f"""
                    SELECT
                        id,
                        original_name,
                        original_extension,
                        mime_type,
                        size_bytes,
                        sha256,
                        source_path,
                        content_json_path,
                        preview_pdf_path,
                        {content_select},
                        status,
                        error_message,
                        chunk_count,
                        used_ocr,
                        created_at,
                        updated_at
                    FROM rag_files
                    ORDER BY updated_at DESC, id DESC
                    """
                )
            ).mappings().fetchall()

        return [_row_to_rag_file(row) for row in rows]

    def list_indexable_files(self) -> list[RagFileRecord]:
        with self._engine.connect() as connection:
            rows = connection.execute(
                text(
                    """
                    SELECT
                        id,
                        original_name,
                        original_extension,
                        mime_type,
                        size_bytes,
                        sha256,
                        source_path,
                        content_json_path,
                        preview_pdf_path,
                        content_json,
                        status,
                        error_message,
                        chunk_count,
                        used_ocr,
                        created_at,
                        updated_at
                    FROM rag_files
                    WHERE status IN ('ready', 'indexing')
                    ORDER BY updated_at DESC, id DESC
                    """
                )
            ).mappings().fetchall()

        return [_row_to_rag_file(row) for row in rows]

    def get_file(
        self,
        file_id: int,
        include_content: bool = True,
    ) -> RagFileRecord | None:
        content_select = "content_json" if include_content else "'' AS content_json"
        with self._engine.connect() as connection:
            row = connection.execute(
                text(
                    f"""
                    SELECT
                        id,
                        original_name,
                        original_extension,
                        mime_type,
                        size_bytes,
                        sha256,
                        source_path,
                        content_json_path,
                        preview_pdf_path,
                        {content_select},
                        status,
                        error_message,
                        chunk_count,
                        used_ocr,
                        created_at,
                        updated_at
                    FROM rag_files
                    WHERE id = :file_id
                    LIMIT 1
                    """
                ),
                {"file_id": file_id},
            ).mappings().fetchone()

        return _row_to_rag_file(row) if row is not None else None

    def get_file_by_sha256(
        self,
        sha256: str,
        include_content: bool = True,
    ) -> RagFileRecord | None:
        content_select = "content_json" if include_content else "'' AS content_json"
        with self._engine.connect() as connection:
            row = connection.execute(
                text(
                    f"""
                    SELECT
                        id,
                        original_name,
                        original_extension,
                        mime_type,
                        size_bytes,
                        sha256,
                        source_path,
                        content_json_path,
                        preview_pdf_path,
                        {content_select},
                        status,
                        error_message,
                        chunk_count,
                        used_ocr,
                        created_at,
                        updated_at
                    FROM rag_files
                    WHERE sha256 = :sha256
                    LIMIT 1
                    """
                ),
                {"sha256": sha256},
            ).mappings().fetchone()

        return _row_to_rag_file(row) if row is not None else None

    def get_preview_pdf_path(self, file_id: int) -> Path:
        record = self.get_file(file_id, include_content=False)
        if record is None:
            raise FileNotFoundError("文件不存在")

        path = _resolve_project_path(record.preview_pdf_path)
        if not path.exists() or not path.is_file():
            raise FileNotFoundError("文件预览尚未生成")

        return path

    def update_chunk_count(self, file_id: int, chunk_count: int) -> None:
        with self._engine.begin() as connection:
            connection.execute(
                text(
                    """
                    UPDATE rag_files
                    SET chunk_count = :chunk_count
                    WHERE id = :file_id
                    """
                ),
                {
                    "file_id": file_id,
                    "chunk_count": max(0, int(chunk_count)),
                },
            )

    def update_file_status(
        self,
        file_id: int,
        status_value: str,
        error_message: str | None = None,
        chunk_count: int | None = None,
        used_ocr: bool | None = None,
    ) -> None:
        assignments = [
            "status = :status",
            "error_message = :error_message",
            "updated_at = :now",
        ]
        payload: dict[str, object] = {
            "file_id": int(file_id),
            "status": status_value,
            "error_message": error_message[:2000] if error_message else None,
            "now": _now_utc(),
        }

        if chunk_count is not None:
            assignments.append("chunk_count = :chunk_count")
            payload["chunk_count"] = max(0, int(chunk_count))

        if used_ocr is not None:
            assignments.append("used_ocr = :used_ocr")
            payload["used_ocr"] = bool(used_ocr)

        with self._engine.begin() as connection:
            connection.execute(
                text(
                    f"""
                    UPDATE rag_files
                    SET {", ".join(assignments)}
                    WHERE id = :file_id
                    """
                ),
                payload,
            )

    def mark_file_ready(self, file_id: int) -> None:
        self.update_file_status(
            file_id,
            RAG_FILE_STATUS_READY,
            error_message=None,
        )

    def mark_file_failed(self, file_id: int, error_message: str) -> None:
        self.update_file_status(
            file_id,
            RAG_FILE_STATUS_FAILED,
            error_message=error_message,
            chunk_count=0,
        )

    def mark_files_indexing(self, file_ids: Iterable[int]) -> None:
        self._update_files_status(
            file_ids,
            RAG_FILE_STATUS_INDEXING,
            error_message=None,
        )

    def mark_files_ready(self, file_ids: Iterable[int]) -> None:
        self._update_files_status(
            file_ids,
            RAG_FILE_STATUS_READY,
            error_message=None,
        )

    def _update_files_status(
        self,
        file_ids: Iterable[int],
        status_value: str,
        error_message: str | None = None,
    ) -> None:
        normalized_ids = sorted({int(file_id) for file_id in file_ids})
        if not normalized_ids:
            return

        placeholders = ", ".join(
            f":file_id_{index}" for index, _ in enumerate(normalized_ids)
        )
        payload: dict[str, object] = {
            "status": status_value,
            "error_message": error_message[:2000] if error_message else None,
            "now": _now_utc(),
        }
        payload.update(
            {
                f"file_id_{index}": file_id
                for index, file_id in enumerate(normalized_ids)
            }
        )

        with self._engine.begin() as connection:
            connection.execute(
                text(
                    f"""
                    UPDATE rag_files
                    SET
                        status = :status,
                        error_message = :error_message,
                        updated_at = :now
                    WHERE id IN ({placeholders})
                    """
                ),
                payload,
            )

    def delete_file(self, file_id: int) -> None:
        record = self.get_file(file_id, include_content=False)
        if record is None:
            raise FileNotFoundError("文件不存在")

        with self._engine.begin() as connection:
            connection.execute(
                text("DELETE FROM rag_files WHERE id = :file_id"),
                {"file_id": file_id},
            )

        for path_value in (
            record.source_path,
            record.content_json_path,
            record.preview_pdf_path,
        ):
            try:
                path = _resolve_project_path(path_value)
            except ValueError:
                continue

            if path.exists() and path.is_file():
                path.unlink()

    @staticmethod
    def _validate_extension(extension: str) -> None:
        if extension not in ALLOWED_RAG_UPLOAD_EXTENSIONS:
            allowed = ", ".join(extension.lstrip(".") for extension in ALLOWED_RAG_UPLOAD_EXTENSIONS)
            raise ValueError(f"仅支持上传以下文件类型: {allowed}")


@lru_cache(maxsize=1)
def get_rag_file_service() -> RagFileService:
    return RagFileService()
